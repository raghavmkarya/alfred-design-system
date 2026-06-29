#!/usr/bin/env python3
"""Copy 'Alfred Deck.docx' and rewrite it with the refined Alfred platform
overview content. Fonts: Bricolage Grotesque (headings) + Plus Jakarta Sans
(body). Keeps the fast-to-edit Word-outline format (Heading per slide + body)."""
import shutil, zipfile, os, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/aryaraghav/Downloads/Alfred Deck.docx"
DST = "/Users/aryaraghav/Downloads/Alfred Deck — Platform Overview.docx"
HEAD = "Bricolage Grotesque"
BODY = "Plus Jakarta Sans"

shutil.copyfile(SRC, DST)
d = Document(DST)

# ---- fonts on styles ----
def set_font(style, name):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)
    for a in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme'):
        if rf.get(qn(a)) is not None:
            del rf.attrib[qn(a)]

by_name = {s.name: s for s in d.styles}
for nm in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
    if nm in by_name: set_font(by_name[nm], HEAD)
for nm in ['normal', 'Normal', 'Subtitle']:
    if nm in by_name: set_font(by_name[nm], BODY)

# document-wide default font (docDefaults) -> body font
styles_el = d.styles.element
dd = styles_el.find(qn('w:docDefaults'))
if dd is not None:
    rprd = dd.find(qn('w:rPrDefault'))
    if rprd is None:
        rprd = OxmlElement('w:rPrDefault'); dd.insert(0, rprd)
    rpr = rprd.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr'); rprd.append(rpr)
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), BODY)
    for a in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme'):
        if rf.get(qn(a)) is not None: del rf.attrib[qn(a)]

# ---- wipe existing body (keep nothing; Word recreates the section) ----
body = d.element.body
for child in list(body):
    body.remove(child)

# ---- helpers ----
def title(text, sub=None):
    p = d.add_paragraph(style='Title'); p.add_run(text)
    if sub:
        sp = d.add_paragraph(style='Subtitle'); sp.add_run(sub)

def slide(n, t):
    p = d.add_paragraph(style='Heading 2')
    r = p.add_run(f"Slide {n} — {t}"); r.bold = True

def lead(text):
    p = d.add_paragraph(style='normal'); r = p.add_run(text); r.bold = True

def para(text):
    p = d.add_paragraph(style='normal'); p.add_run(text)

def item(label, rest):
    p = d.add_paragraph(style='normal')
    r = p.add_run(label); r.bold = True
    p.add_run(" " + rest)

def gap():
    d.add_paragraph(style='normal')

# ---- content ----
title("Alfred — Platform Overview",
      "The AI memory powering every decision across your organisation.")
gap()

slide(1, "Cover")
para("Decision intelligence platform")
lead("The AI memory powering every decision across your organisation.")
para("Alfred connects your entire business stack, reasons across it continuously, and tells every "
     "leader what changed, why, and what to do next.")
gap()

slide(2, "The Gap")
lead("You have the data. You don't have the decision.")
item("Fragmented across the stack.", "Every function runs 5–15 systems that don't talk to each other in plain language.")
item("Dashboards show symptoms.", "A number moved — but not whether the cause is fatigue, an algorithm, a tracking error, or a competitor.")
item("The answer lands too late.", "By the time an analyst reconciles it, the budget — and the moment — has already drifted.")
para("54% of leaders say connecting data sources is a major barrier to insight (NIQ 2026). "
     "63% miss opportunities because decisions take too long (PwC 2025). One in three rely on "
     "5–15 tools just to measure ROI (NIQ 2026).")
gap()

slide(3, "The Cost of Slow")
lead("The cost of slow.")
para("Every hour between issue and action, the cost compounds against you.")
para("Speed of correction now beats quality of launch — in every function.")
gap()

slide(4, "What Alfred Is")
lead("Not another dashboard. A decision intelligence platform that tells every leader what changed, "
     "why, and what to do — continuously, across the whole business.")
para("Alfred sits between your business data and your leadership decisions. It connects the stack, "
     "reasons across it 24/7, and turns hundreds of signals into the handful of decisions that "
     "actually need you.")
para("It doesn't replace your tools — it makes sense of them.")
gap()

slide(5, "Three Properties")
lead("Everyone else describes the problem. Alfred decides.")
item("Proactive, not reactive.", "BI waits for you to ask. Alfred surfaces the issue, risk, or opening before you go looking — the first sign of trouble is an Alfred alert, not a board question.")
item("Explains, rather than reports.", "“CPA up 34%” is a notification. “CPA up 34% from creative fatigue in 25–34 — refresh, or move 15% to LinkedIn” is intelligence.")
item("Insight to execution, in one step.", "Approve a recommendation and Alfred acts inside your connected tools — every change logged and traceable.")
gap()

slide(6, "How It Works")
lead("Four stages, fully automated.")
item("Connect.", "Read-only API links to your existing stack. Up to a year of history synced at onboarding. Zero engineering.")
item("Analyze.", "Continuous correlation across every channel, system, and funnel stage — 24 hours a day.")
item("Recommend.", "Decision-ready intelligence in plain language, prioritised by business impact, with the reasoning attached.")
item("Act.", "Approve and Alfred executes inside your tools. Every change recorded in a full audit trail.")
gap()

slide(7, "What a Leader Receives")
lead("Three outputs. Every day. In plain language.")
item("A prioritised daily brief.", "What changed overnight, what caused it, and what to act on today — ranked by business impact, ready before your first meeting.")
item("On-demand answers.", "Ask any business question in plain language and get a data-grounded answer in seconds. Seek Alfred — no analyst, no waiting.")
item("Proactive alerts.", "Risks and opportunities surfaced before you go looking — each one explained, prioritised, and paired with a recommended action.")
gap()

slide(8, "Alfred Core")
lead("The organisational brain underneath every module.")
para("Each module runs a network of specialised agents. Alfred Core is the layer that connects them — "
     "routing every signal, tracing root causes across functions, and keeping the whole leadership "
     "team aligned to one view of reality.")
item("Cross-function intelligence.", "A marketing signal with a sales implication, a revenue risk with a finance one — connected, not siloed.")
item("Root-cause intelligence.", "Every tool today is diagnostic. Alfred is etiological — it traces the full causal chain to the one cause.")
item("Institutional knowledge.", "The context that normally walks out the door when a key person leaves. Alfred keeps it, permanently.")
item("Compounding intelligence.", "Every outcome recorded makes every future recommendation more accurate — no retraining required.")
gap()

slide(9, "One Root Cause, One Fix")
lead("Four symptoms. One root cause. One fix.")
para("Operations cut resourcing six weeks ago to protect margin → two late-stage deals stall → a "
     "$380k forecast downgrade lands in the pipeline → a Q3 marketing budget increase is paused in response.")
para("Every other tool sees four separate problems and proposes four separate fixes. Alfred Core "
     "traces the chain across functions and surfaces the single cause underneath — with a confidence "
     "score on every link.")
gap()

slide(10, "The Platform: Five Modules")
lead("Five modules, one intelligence layer.")
item("Marketing — Live.", "Spend, creative & brand-visibility intelligence.")
item("Sales — In development.", "Pipeline health, deal patterns & revenue foresight.")
item("Finance — Planned.", "Spend efficiency, forecasting & cost intelligence.")
item("Operations — Planned.", "Capacity, process health & cross-functional coordination.")
item("Founders — Planned.", "Whole-business view, strategy & risk synthesis.")
para("Every module is its own agent network — all running on Alfred Core, sharing one memory and one "
     "signal router. The platform gets more valuable with every function added.")
gap()

slide(11, "The Live Wedge — Proof")
lead("What changes when Alfred runs a leader's mornings.")
para("70% faster decision cycles. 12% less wasted spend. 15+ hours reclaimed each week. 85% of "
     "questions answered without an analyst.")
para("From Alfred for Marketing, the first live module. Time to first insight: under a week. Every "
     "additional module compounds on the same engine and the same memory.")
gap()

slide(12, "The Moat: Memory")
lead("It gets sharper every month — and competitors can't backfill that.")
item("Institutional memory.", "Alfred Core learns this company's patterns, outcomes, and causes — not industry averages.")
item("Two-way integrations.", "It doesn't just read — it acts, and records what worked.")
item("Compounding accuracy.", "Recommendations get more precise without retraining a single model.")
para("Six months in, the switching cost isn't the product — it's the institutional memory accumulated inside it.")
gap()

slide(13, "Why Alfred Wins")
lead("Built for the person at the top — not the data team.")
item("Versus BI tools.", "Built for the data team to explore and report. Alfred is built for the leader — decisions, not queries.")
item("Versus generic AI.", "Helps you work faster in your documents. Alfred helps you decide faster, in your data.")
item("Versus point analytics.", "Tells you which channel or deal worked. Alfred tells you what it means for next quarter.")
item("Versus doing nothing.", "Manual reconciliation; the cost compounds. Alfred closes the issue-to-action gap to under 2 hours.")
para("Dashboards show. Copilots wait. Analysts leave. None of them remember. Alfred does.")
gap()

slide(14, "Integrations & Trust")
lead("Connects to the stack you already run. No rip and replace.")
para("Google Ads, Meta Ads, LinkedIn Ads, Amazon Ads, Google Analytics 4, Google Search Console, "
     "HubSpot, Salesforce, Slack, Email.")
para("All connections are read-only by default. Up to a year of history is synced at onboarding. "
     "Alfred only writes back inside a connected tool when you explicitly approve it — and every "
     "change is recorded in a full audit trail. No cross-customer training; your intelligence stays yours.")
gap()

slide(15, "Get Started")
lead("From signup to first brief in under a week.")
item("Week 1 — Connect.", "Secure, read-only integrations. Historical data synced. Zero engineering required.")
item("Weeks 1–2 — Calibrate.", "KPIs, alert thresholds, channels, and objectives configured to your business.")
item("Week 2 — Brief.", "First intelligence briefing delivered. Fully operational.")
gap()

slide(16, "Closing")
lead("Build the enterprise brain. One function at a time.")
para("See what your business looks like with one memory instead of five — and be the most prepared "
     "leadership team in your market, every morning.")
para("hello@seekalfred.ai · seekalfred.ai")

d.save(DST)

# ---- patch theme fonts (major=headings, minor=body) so any inherited default matches ----
tmp = DST + ".tmp"
with zipfile.ZipFile(DST, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
    for it in zin.namelist():
        data = zin.read(it)
        if it == "word/theme/theme1.xml":
            t = data.decode('utf-8')
            t = re.sub(r'(<a:majorFont>\s*<a:latin[^>]*typeface=")[^"]*(")',
                       r'\g<1>' + HEAD + r'\g<2>', t, count=1)
            t = re.sub(r'(<a:minorFont>\s*<a:latin[^>]*typeface=")[^"]*(")',
                       r'\g<1>' + BODY + r'\g<2>', t, count=1)
            data = t.encode('utf-8')
        zout.writestr(it, data)
os.replace(tmp, DST)

doc2 = Document(DST)
print("OK:", DST)
print("paragraphs:", len(doc2.paragraphs))
print("Heading2 font:", by_name['Heading 2'].font.name, "| normal font:", by_name['normal'].font.name)
